import os
import pandas as pd
import re
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FileProcessor:

    SUPPORTED_FORMATS = ['.pdf', '.docx', '.xlsx', '.xls', '.csv', '.txt']

    @staticmethod
    def clean_text(text: str) -> str:
        if not text or len(text.strip()) == 0:
            return ""

        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\s*\n\s*', '\n', text)

        noise_patterns = [
            r'Page \d+ of \d+',
            r'© \d{4}',
            r'Confidential',
            r'www\.[\w\-\./]+',
            r'https?://\S+',
        ]
        for pattern in noise_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)

        lines = text.split('\n')
        cleaned_lines = [line.strip() for line in lines if len(line.strip()) >= 3]
        result = ' '.join(cleaned_lines)
        logger.info(f"Cleaned text length: {len(result)}")
        return result.strip()

    @staticmethod
    def _extract_pdf_text(file_path: str) -> str:
        """
        Strategy:
        1. Try pdfplumber (best for complex/formatted PDFs).
        2. Fallback to pypdf.
        3. If both yield < 200 chars, attempt OCR via pdf2image + pytesseract.
        """

        # ── Strategy 1: pdfplumber ──────────────────────────────────────────
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages[:150]:
                    txt = page.extract_text()
                    if txt and len(txt.strip()) > 5:
                        text_parts.append(txt)
                    # Also extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            row_text = " | ".join(
                                str(cell).strip() for cell in row if cell and str(cell).strip()
                            )
                            if row_text:
                                text_parts.append(row_text)

            plumber_text = "\n".join(text_parts)
            logger.info(f"pdfplumber extracted {len(plumber_text)} chars from {os.path.basename(file_path)}")

            if len(plumber_text.strip()) >= 150:
                return FileProcessor.clean_text(plumber_text)
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {e}")
            plumber_text = ""

        # ── Strategy 2: pypdf ───────────────────────────────────────────────
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages[:150]:
                txt = page.extract_text()
                if txt and len(txt.strip()) > 5:
                    text_parts.append(txt)
            pypdf_text = "\n".join(text_parts)
            logger.info(f"pypdf extracted {len(pypdf_text)} chars from {os.path.basename(file_path)}")

            if len(pypdf_text.strip()) >= 150:
                return FileProcessor.clean_text(pypdf_text)
        except Exception as e:
            logger.warning(f"pypdf failed for {file_path}: {e}")
            pypdf_text = ""

        # ── Strategy 3: OCR fallback ────────────────────────────────────────
        logger.info(f"Both extractors returned short text — attempting OCR for {os.path.basename(file_path)}")
        return FileProcessor._ocr_pdf(file_path)

    @staticmethod
    def _ocr_pdf(file_path: str) -> str:
        """Convert PDF pages to images and run pytesseract OCR on each."""
        try:
            from pdf2image import convert_from_path
            import pytesseract
            # On Windows, set the tesseract path if needed:
            # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        except ImportError as e:
            logger.error(f"OCR dependencies missing: {e}. Run: pip install pdf2image pytesseract")
            return ""

        try:
            images = convert_from_path(file_path, dpi=250, first_page=1, last_page=50)
            logger.info(f"OCR: converted {len(images)} pages to images")

            ocr_texts = []
            for i, img in enumerate(images):
                try:
                    text = pytesseract.image_to_string(img, lang="eng")
                    if text and len(text.strip()) > 10:
                        ocr_texts.append(text)
                        logger.info(f"OCR page {i+1}: {len(text)} chars")
                except Exception as page_err:
                    logger.warning(f"OCR failed on page {i+1}: {page_err}")

            full_text = "\n".join(ocr_texts)
            logger.info(f"OCR total: {len(full_text)} chars from {len(ocr_texts)} pages")

            if len(full_text.strip()) < 50:
                logger.error("OCR produced too little text — PDF may be corrupt or purely graphical.")
                return ""

            return FileProcessor.clean_text(full_text)

        except Exception as e:
            logger.error(f"OCR pipeline failed for {file_path}: {e}", exc_info=True)
            return ""

    @staticmethod
    def extract_text(file_path: str) -> str:
        if not os.path.exists(file_path):
            logger.warning(f"File not found: {file_path}")
            return ""

        ext = os.path.splitext(file_path)[1].lower()

        if ext not in FileProcessor.SUPPORTED_FORMATS:
            logger.warning(f"Unsupported file extension: {ext}")
            return ""

        try:
            if ext == '.pdf':
                return FileProcessor._extract_pdf_text(file_path)

            elif ext == '.docx':
                import docx
                doc = docx.Document(file_path)
                paragraphs = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        paragraphs.append(p.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(
                            cell.text.strip() for cell in row.cells if cell.text.strip()
                        )
                        if row_text:
                            paragraphs.append(row_text)
                raw = "\n".join(paragraphs)
                logger.info(f"DOCX raw extraction: {len(raw)} chars")
                return FileProcessor.clean_text(raw)

            elif ext in ['.xlsx', '.xls']:
                xl = pd.ExcelFile(file_path)
                all_text = []
                for sheet_name in xl.sheet_names:
                    df = xl.parse(sheet_name)
                    sheet_text = f"Sheet: {sheet_name} | {df.shape[0]} rows x {df.shape[1]} columns\n"
                    sheet_text += df.head(100).to_string()
                    all_text.append(sheet_text)
                return "\n\n".join(all_text)

            elif ext == '.csv':
                encodings = ["utf-8", "latin1", "cp1252"]
                for enc in encodings:
                    try:
                        df = pd.read_csv(file_path, encoding=enc)
                        text = f"CSV dataset with {df.shape[0]} rows and {df.shape[1]} columns.\n"
                        text += df.head(100).to_string()
                        logger.info(f"CSV read OK with encoding {enc}: {len(text)} chars")
                        return text
                    except Exception as csv_err:
                        logger.warning(f"CSV read failed with encoding {enc}: {csv_err}")
                logger.error(f"All encoding attempts failed for CSV: {file_path}")
                return ""

            elif ext == '.txt':
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                logger.info(f"TXT raw extraction: {len(text)} chars")
                return FileProcessor.clean_text(text)

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}", exc_info=True)

        return ""